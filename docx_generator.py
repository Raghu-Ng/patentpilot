from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
from models import Drawing

class PatentDocxGenerator:
    """Generate patent specification DOCX files"""
    
    def __init__(self, template_path="patent_specification_template.docx"):
        """Initialize with template path"""
        self.template_path = template_path
        self.output_dir = "generated_docs"
        
        # Create output directory if it doesn't exist
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def generate_patent_docx(self, draft):
        """Generate complete patent specification DOCX"""
        try:
            # Load template
            doc = DocxTemplate(self.template_path)
            
            # Get drawings for this draft
            drawings = Drawing.objects(draft_id=str(draft.id))
            
            # Prepare context data
            context = self._prepare_context(draft, drawings)
            
            # Render template
            doc.render(context)
            
            # Generate output filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"patent_specification_{draft.id}_{timestamp}.docx"
            output_path = os.path.join(self.output_dir, output_filename)
            
            # Save document
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error generating DOCX: {str(e)}")
    
    def _prepare_context(self, draft, drawings):
        """Prepare context data for template rendering"""
        
        # Basic information
        context = {
            'title': draft.title or 'Patent Specification',
            'field_of_invention': draft.field_of_invention or '',
            'brief_summary': draft.brief_summary or '',
            'key_components': draft.key_components or '',
            'problem_solved': draft.problem_solved or '',
            
            # Patent sections
            'background': draft.background or '',
            'summary': draft.summary or '',
            'detailed_description': draft.detailed_description or '',
            'claims': draft.claims or '',
            'abstract': draft.abstract or '',
            
            # Metadata
            'generation_date': datetime.now().strftime('%B %d, %Y'),
            'draft_id': str(draft.id),
            'current_step': draft.current_step,
            'is_complete': draft.is_complete,
            
            # Drawings
            'drawings': [],
            'has_drawings': False
        }
        
        # Process drawings
        if drawings:
            context['has_drawings'] = True
            for i, drawing in enumerate(drawings, 1):
                try:
                    # Create inline image for DOCX
                    if os.path.exists(drawing.file_path):
                        img = InlineImage(
                            doc, 
                            drawing.file_path, 
                            width=Mm(120)  # Adjust size as needed
                        )
                        context['drawings'].append({
                            'number': i,
                            'filename': drawing.original_filename,
                            'description': drawing.description or f'Figure {i}',
                            'image': img
                        })
                except Exception as e:
                    # If image processing fails, add text reference
                    context['drawings'].append({
                        'number': i,
                        'filename': drawing.original_filename,
                        'description': drawing.description or f'Figure {i}',
                        'image': None,
                        'error': str(e)
                    })
        
        return context
    
    def create_template(self):
        """Create a basic patent specification template if none exists"""
        if os.path.exists(self.template_path):
            return
        
        # Create a basic template structure
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('PATENT SPECIFICATION', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title of Invention
        doc.add_heading('Title of Invention', level=1)
        doc.add_paragraph('{{ title }}')
        
        # Field of Invention
        doc.add_heading('Field of Invention', level=1)
        doc.add_paragraph('{{ field_of_invention }}')
        
        # Background
        doc.add_heading('Background of Invention', level=1)
        doc.add_paragraph('{{ background }}')
        
        # Summary
        doc.add_heading('Summary of Invention', level=1)
        doc.add_paragraph('{{ summary }}')
        
        # Detailed Description
        doc.add_heading('Detailed Description', level=1)
        doc.add_paragraph('{{ detailed_description }}')
        
        # Claims
        doc.add_heading('Claims', level=1)
        doc.add_paragraph('{{ claims }}')
        
        # Abstract
        doc.add_heading('Abstract', level=1)
        doc.add_paragraph('{{ abstract }}')
        
        # Drawings section
        if '{{ has_drawings }}':
            doc.add_heading('Drawings', level=1)
            for drawing in '{{ drawings }}':
                doc.add_paragraph(f"Figure {drawing.number}: {drawing.description}")
                if drawing.image:
                    doc.add_picture(drawing.image)
        
        # Save template
        doc.save(self.template_path)

def generate_patent_docx(draft):
    """Convenience function to generate patent DOCX"""
    generator = PatentDocxGenerator()
    
    # Create template if it doesn't exist
    generator.create_template()
    
    # Generate the document
    return generator.generate_patent_docx(draft) 