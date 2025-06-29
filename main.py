from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

app = Flask(__name__)

def calculate_fees(form_data):
    # Base fees
    fees = {
        'filing_fee': 0,
        'publication_fee': 0,
        'examination_fee': 0
    }
    
    # Filing fee calculation
    if form_data['applicationType'] == 'Provisional':
        fees['filing_fee'] = 1600  # Base provisional filing fee
    else:  # Complete application
        fees['filing_fee'] = 8000  # Base complete filing fee
        
        # Additional fees for early publication
        if form_data.get('publicationPreference') == 'Early':
            fees['publication_fee'] = 2500
        else:
            fees['publication_fee'] = 0
            
        # Additional fees for expedited examination
        if form_data.get('examinationPreference') == 'Expedited':
            fees['examination_fee'] = 12000
        else:
            fees['examination_fee'] = 4000
    
    # Apply category-based discounts
    if form_data.get('preConfigureApplicant') == 'Yes':
        category = form_data.get('preConfiguredApplicant', {}).get('category')
        if category == 'Startup':
            fees['filing_fee'] *= 0.8  # 20% discount
            fees['publication_fee'] *= 0.8
            fees['examination_fee'] *= 0.8
        elif category == 'Small':
            fees['filing_fee'] *= 0.9  # 10% discount
            fees['publication_fee'] *= 0.9
            fees['examination_fee'] *= 0.9
    
    return fees

def generate_document(form_data):
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('PATENT APPLICATION FORM', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Application Type
    doc.add_heading('Application Type', level=2)
    doc.add_paragraph(f"Type: {form_data['applicationType']}")
    
    if form_data['applicationType'] == 'Complete':
        if form_data.get('previousProvisionalFiled') == 'Yes':
            doc.add_paragraph(f"Previous Provisional Application Number: {form_data.get('provisionalApplicationNumber', 'N/A')}")
    
    # Applicants Section
    doc.add_heading('3A. APPLICANTS', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Name in Full'
    header_cells[1].text = 'Nationality'
    header_cells[2].text = 'Country of Residence'
    header_cells[3].text = 'Address of the Applicant'
    
    # Add applicant rows
    for applicant in form_data.get('applicants', []):
        row_cells = table.add_row().cells
        row_cells[0].text = applicant.get('name', 'None')
        row_cells[1].text = applicant.get('nationality', 'None')
        row_cells[2].text = applicant.get('residency', 'None')
        row_cells[3].text = applicant.get('address', 'None')
    
    # Add empty rows if less than 3 applicants
    while len(table.rows) < 4:  # Header + 3 applicants
        row_cells = table.add_row().cells
        for cell in row_cells:
            cell.text = 'None'
    
    # Category of Applicant
    doc.add_heading('3B. CATEGORY OF APPLICANT', level=2)
    if form_data.get('preConfigureApplicant') == 'Yes':
        category = form_data.get('preConfiguredApplicant', {}).get('category', '')
        categories = ['Natural Person', 'Other Than Natural Person', 'Educational institution', 
                     'Small Entity', 'Start-Up', 'Others']
        for cat in categories:
            p = doc.add_paragraph()
            p.add_run('☒ ' if cat == category else '☐ ')
            p.add_run(cat)
    
    # Inventors Section
    doc.add_heading('4. INVENTORS', level=2)
    same_as_applicants = form_data.get('sameAsApplicants', 'No')
    p = doc.add_paragraph()
    p.add_run('Are/Is all the inventors same as the applicants named above? ').bold = True
    p.add_run('☒ ' if same_as_applicants == 'Yes' else '☐ ')
    p.add_run('Yes ')
    p.add_run('☐ ' if same_as_applicants == 'Yes' else '☒ ')
    p.add_run('No')
    
    if same_as_applicants == 'No':
        doc.add_paragraph('If "No", furnish the details of the inventors')
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Name in Full'
        header_cells[1].text = 'Nationality'
        header_cells[2].text = 'Country of Residence'
        header_cells[3].text = 'Address of inventor'
        
        # Add inventor rows
        for inventor in form_data.get('inventors', []):
            row_cells = table.add_row().cells
            row_cells[0].text = inventor.get('name', 'None')
            row_cells[1].text = inventor.get('nationality', 'None')
            row_cells[2].text = inventor.get('residency', 'None')
            row_cells[3].text = inventor.get('address', 'None')
        
        # Add empty rows if less than 1 inventor
        while len(table.rows) < 2:  # Header + 1 inventor
            row_cells = table.add_row().cells
            for cell in row_cells:
                cell.text = 'None'
    
    # Title of Invention
    doc.add_heading('5. TITLE OF THE INVENTION', level=2)
    doc.add_paragraph(form_data.get('title', ''))
    
    # Agent Details
    doc.add_heading('6. AUTHORISED REGISTERED PATENT AGENTS', level=2)
    agent = form_data.get('agent', {})
    p = doc.add_paragraph()
    p.add_run(f"IN/PA No: {agent.get('inpaNo', '')}\n")
    p.add_run(f"Name: {agent.get('agentName', '')}\n")
    p.add_run(f"Mobile number: {agent.get('agentMobile', '')}")
    
    # Save the document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"patent_application_{timestamp}.docx"
    doc.save(filename)
    return filename

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/submit', methods=['POST'])
def submit():
    form_data = request.form.to_dict()
    
    # Process applicants data
    applicants = []
    i = 0
    while f'applicants[{i}][name]' in form_data:
        applicant = {
            'name': form_data[f'applicants[{i}][name]'],
            'nationality': form_data[f'applicants[{i}][nationality]'],
            'residency': form_data[f'applicants[{i}][residency]'],
            'address': form_data[f'applicants[{i}][address]']
        }
        applicants.append(applicant)
        i += 1
    form_data['applicants'] = applicants
    
    # Process inventors data
    inventors = []
    i = 0
    while f'inventors[{i}][name]' in form_data:
        inventor = {
            'name': form_data[f'inventors[{i}][name]'],
            'gender': form_data[f'inventors[{i}][gender]'],
            'nationality': form_data[f'inventors[{i}][nationality]'],
            'residency': form_data[f'inventors[{i}][residency]'],
            'address': form_data[f'inventors[{i}][address]']
        }
        if inventor['residency'] == 'India':
            inventor['state'] = form_data.get(f'inventors[{i}][state]', '')
        inventors.append(inventor)
        i += 1
    form_data['inventors'] = inventors
    
    # Process sheet counts
    sheet_counts = {
        'patentDocumentSheets': int(form_data.get('sheetCounts[patentDocumentSheets]', 0)),
        'abstractSheets': int(form_data.get('sheetCounts[abstractSheets]', 0)),
        'claimsSheets': int(form_data.get('sheetCounts[claimsSheets]', 0)),
        'drawingSheets': int(form_data.get('sheetCounts[drawingSheets]', 0))
    }
    form_data['sheetCounts'] = sheet_counts
    
    # Process agent details
    agent = {
        'inpaNo': form_data.get('agent[inpaNo]', ''),
        'agentName': form_data.get('agent[agentName]', ''),
        'agentMobile': form_data.get('agent[agentMobile]', ''),
        'agentEmail': form_data.get('agent[agentEmail]', '')
    }
    form_data['agent'] = agent
    
    # Process service address
    service_address = {
        'serviceName': form_data.get('serviceAddress[serviceName]', ''),
        'postalAddress': form_data.get('serviceAddress[postalAddress]', ''),
        'telephone': form_data.get('serviceAddress[telephone]', ''),
        'mobile': form_data.get('serviceAddress[mobile]', ''),
        'fax': form_data.get('serviceAddress[fax]', ''),
        'email': form_data.get('serviceAddress[email]', '')
    }
    form_data['serviceAddress'] = service_address
    
    # Process pre-configured applicant if present
    if form_data.get('preConfigureApplicant') == 'Yes':
        pre_configured = {
            'name': form_data.get('preConfiguredApplicant[name]', ''),
            'category': form_data.get('preConfiguredApplicant[category]', '')
        }
        form_data['preConfiguredApplicant'] = pre_configured
    
    # Generate the document
    filename = generate_document(form_data)
    
    # Send the file to the user
    return send_file(filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
